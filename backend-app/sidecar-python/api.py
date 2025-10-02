#!/usr/bin/env python3
"""
API FastAPI locale pour génération de documents Word
Intégrée dans l'application Tauri comme sidecar
"""

import os
import sys
import logging
from datetime import datetime
from typing import List, Optional

import uvicorn
from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document

# Configuration - LOCALHOST ONLY
API_HOST = "127.0.0.1"
API_PORT = 8001
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Proposal Document Generator", version="1.0.0")

# CORS pour Tauri uniquement
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:*", "http://127.0.0.1:*", "tauri://localhost"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Models
class Product(BaseModel):
    product_type: str
    user_count: int
    standalone_count: int = 0
    server_key_count: int = 0
    unit_price: float = 0
    total_price: float = 0
    annual_reduction: float = 0
    licence: bool = False
    training: bool = False
    training_days: int = 0
    training_cost_per_day: float = 1500
    support: bool = False
    support_years: int = 0

class Contact(BaseModel):
    first_name: str
    last_name: str
    email: str
    phone_number: Optional[str] = None

class Company(BaseModel):
    name: str
    address: Optional[str] = None
    city: Optional[str] = None
    postal_code: Optional[str] = None
    country: Optional[str] = None

class ProposalRequest(BaseModel):
    proposal_number: str
    company: Company
    contact: Optional[Contact] = None
    products: List[Product]
    currency: str = "USD"
    valid_until: Optional[str] = None
    notes: Optional[str] = None

def format_validity_date(valid_until):
    """Formate la date de validité"""
    if not valid_until:
        return "30 days"
    
    try:
        if isinstance(valid_until, str) and "/" in valid_until:
            return valid_until
        
        from datetime import datetime
        if isinstance(valid_until, str):
            for fmt in ["%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%d/%m/%Y"]:
                try:
                    date_obj = datetime.strptime(valid_until.split('T')[0], "%Y-%m-%d")
                    return date_obj.strftime("%d/%m/%Y")
                except ValueError:
                    continue
        
        return valid_until if valid_until else "30 days"
    except Exception:
        return "30 days"

def generate_licence_description(product):
    """Génère description des licences"""
    if not product or not product.get('licence', False):
        return "0"
    
    standalone_count = product.get('standalone_count', 0) or 0
    server_key_count = product.get('server_key_count', 0) or 0
    
    parts = []
    if standalone_count > 0:
        parts.append(f"{standalone_count} standalone")
    if server_key_count > 0:
        parts.append(f"{server_key_count} server key{'s' if server_key_count > 1 else ''}")
    
    if not parts:
        return f"{product.get('user_count', 0)} user{'s' if product.get('user_count', 0) > 1 else ''}"
    
    return " and ".join(parts)

def remove_table_with_immediate_title(doc, table_element):
    """Supprime un tableau et son titre précédent"""
    try:
        parent = table_element.getparent()
        table_position = list(parent).index(table_element)
        elements_to_remove = [table_element]
        
        current_position = table_position - 1
        while current_position >= 0:
            current_element = parent[current_position]
            if current_element.tag.endswith('p'):
                element_text = ""
                for para in doc.paragraphs:
                    if para._element == current_element:
                        element_text = para.text.strip()
                        break
                
                if len(element_text) == 0:
                    elements_to_remove.insert(0, current_element)
                else:
                    elements_to_remove.insert(0, current_element)
                    break
            elif current_element.tag.endswith('tbl'):
                break
            else:
                elements_to_remove.insert(0, current_element)
            current_position -= 1
        
        for element in elements_to_remove:
            element.getparent().remove(element)
        return True
    except Exception as e:
        logger.error(f"Erreur suppression table: {e}")
        try:
            table_element.getparent().remove(table_element)
            return True
        except:
            return False

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Remplace texte dans un paragraphe en conservant le formatage"""
    if old_text in paragraph.text:
        full_text = paragraph.text.replace(old_text, new_text)
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)
        return True
    return False

def process_word_document(doc, template_data, product=None):
    """Traite le document Word avec les remplacements"""
    try:
        replacements_made = 0
        
        # Remplacer dans paragraphes
        for paragraph in doc.paragraphs:
            for key, value in template_data.items():
                if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                    replacements_made += 1
        
        # Remplacer dans tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in template_data.items():
                            if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                replacements_made += 1
        
        # Remplacer dans headers/footers
        # Remplacement FORCÉ dans tous les paragraphes
        for paragraph in doc.paragraphs:
            for key, value in template_data.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
                    replacements_made += 1
                    logger.info(f"Remplacement FORCÉ: {{{key}}} -> {value}")
        
        # Traitement spécial page de garde
        logger.info("=== TRAITEMENT PAGE DE GARDE ===")
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if any(f"{{{key}}}" in paragraph_text for key in template_data.keys()):
                logger.info(f"Paragraphe page de garde: '{paragraph_text[:50]}...'")
                for key, value in template_data.items():
                    if f"{{{key}}}" in paragraph_text:
                        paragraph.text = paragraph_text.replace(f"{{{key}}}", str(value))
                        replacements_made += 1
                        logger.info(f"Remplacement PAGE DE GARDE: {{{key}}} -> {value}")
        
        # Traitement des en-têtes de première page (si différent)
        logger.info("=== TRAITEMENT EN-TÊTES PREMIÈRE PAGE ===")
        for section in doc.sections:
            if section.first_page_header:
                logger.info("En-tête première page trouvé")
                for paragraph in section.first_page_header.paragraphs:
                    for key, value in template_data.items():
                        if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                            replacements_made += 1
                            logger.info(f"Remplacement EN-TÊTE PREMIÈRE PAGE: {{{key}}} -> {value}")
                for table in section.first_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in template_data.items():
                                    if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                        replacements_made += 1
                                        logger.info(f"Remplacement EN-TÊTE PREMIÈRE PAGE (tableau): {{{key}}} -> {value}")
            
            # Pieds de page de première page
            if section.first_page_footer:
                logger.info("Pied de page première page trouvé")
                for paragraph in section.first_page_footer.paragraphs:
                    for key, value in template_data.items():
                        if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                            replacements_made += 1
                            logger.info(f"Remplacement PIED PREMIÈRE PAGE: {{{key}}} -> {value}")
                for table in section.first_page_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in template_data.items():
                                    if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                        replacements_made += 1
                                        logger.info(f"Remplacement PIED PREMIÈRE PAGE (tableau): {{{key}}} -> {value}")
            
            # En-têtes de pages paires (si différent) - avec gestion d'erreur
            try:
                if hasattr(section, 'even_page_header') and section.even_page_header:
                    logger.info("En-tête pages paires trouvé")
                    for paragraph in section.even_page_header.paragraphs:
                        for key, value in template_data.items():
                            if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                replacements_made += 1
                                logger.info(f"Remplacement EN-TÊTE PAGES PAIRES: {{{key}}} -> {value}")
                    for table in section.even_page_header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for key, value in template_data.items():
                                        if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                            replacements_made += 1
                                            logger.info(f"Remplacement EN-TÊTE PAGES PAIRES (tableau): {{{key}}} -> {value}")
            except Exception as e:
                logger.info(f"En-tête pages paires non disponible ou erreur: {e}")
            
            # En-têtes et pieds normaux
            if section.header:
                for paragraph in section.header.paragraphs:
                    for key, value in template_data.items():
                        if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                            replacements_made += 1
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in template_data.items():
                                    if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                                        replacements_made += 1
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for key, value in template_data.items():
                        if replace_text_in_paragraph(paragraph, f"{{{key}}}", str(value)):
                            replacements_made += 1
        
        
        # Logique suppression tableaux selon produit
        if product:
            product_type = product.get('product_type', '')
            is_licence = product.get('licence', False)
            is_support = product.get('support', False)
            annual_reduction = product.get('annual_reduction', 0)
            
            tables_to_keep = [0]  # Toujours garder tableau 0
            
            # Tableaux de licence selon produit
            if "HTZ Communications" in product_type:
                tables_to_keep.append(2 if annual_reduction > 0 else 1)
            elif "HTZ Warfare" in product_type:
                tables_to_keep.append(4 if annual_reduction > 0 else 3)
            elif "ICS Manager" in product_type:
                tables_to_keep.append(6 if annual_reduction > 0 else 5)
            
            # Tableau training
            if product.get('training', False):
                tables_to_keep.append(7)
            
            # Tableau support
            if is_support:
                if "HTZ Communications" in product_type:
                    tables_to_keep.append(10)
                elif "HTZ Warfare" in product_type:
                    tables_to_keep.append(9)
                elif "ICS Manager" in product_type:
                    tables_to_keep.append(12)
            
            # Supprimer les autres tableaux
            num_tables = len(doc.tables)
            tables_to_remove = [i for i in range(num_tables) if i not in tables_to_keep]
            
            for table_index in sorted(tables_to_remove, reverse=True):
                if table_index < len(doc.tables):
                    remove_table_with_immediate_title(doc, doc.tables[table_index]._element)
        
        logger.info(f"Remplacements: {replacements_made}")
        return replacements_made
    except Exception as e:
        logger.error(f"Erreur traitement: {e}")
        return 0

@app.post("/generate-word")
async def generate_word(request: ProposalRequest):
    """Génère un document Word à partir du template"""
    try:
        logger.info(f"Génération Word pour: {request.company.name}")
        
        if not os.path.exists(TEMPLATE_PATH):
            raise HTTPException(status_code=404, detail="Template not found")
        
        doc = Document(TEMPLATE_PATH)
        
        # Calculer total
        total_amount = sum(p.total_price for p in request.products)
        first_product = request.products[0] if request.products else None
        
        # Template data
        template_data = {
            "companyname": request.company.name,
            "companyadresse": request.company.address or "",
            "villename": request.company.city or "",
            "zipcode": request.company.postal_code or "",
            "offerreference": request.proposal_number,
            "offerdate": datetime.now().strftime("%d/%m/%Y"),
            "offervalidity": format_validity_date(request.valid_until),
            "pricetotal": f"{total_amount:,.2f}",
            "pricetotalsupport": f"{total_amount:,.2f}",
            "notes": request.notes or f"Proposal for {request.company.name}",
            "name": request.contact.first_name if request.contact else "",
            "lastname": request.contact.last_name if request.contact else "",
            "email": request.contact.email if request.contact else "",
            "phone": request.contact.phone_number if request.contact else "",
        }
        
        if first_product:
            product_dict = {
                'product_type': first_product.product_type,
                'user_count': first_product.user_count,
                'standalone_count': first_product.standalone_count,
                'server_key_count': first_product.server_key_count,
                'unit_price': first_product.unit_price,
                'licence': first_product.licence,
                'training': first_product.training,
                'training_days': first_product.training_days,
                'training_cost_per_day': first_product.training_cost_per_day,
                'support': first_product.support,
                'support_years': first_product.support_years,
                'annual_reduction': first_product.annual_reduction,
            }
            
            template_data.update({
                "product": first_product.product_type,
                "user": str(first_product.user_count),
                "licence": generate_licence_description(product_dict),
                "costslicence": f"{first_product.unit_price:,.2f}",
                "discount": f"{first_product.annual_reduction}%",
                "trainingprice": f"{first_product.training_days * first_product.training_cost_per_day:,.2f}" if first_product.training else "0",
                "traningday": str(first_product.training_days),
                "years": str(first_product.support_years),
            })
            
            # Support price
            if first_product.support:
                base_price = first_product.unit_price * (1 - first_product.annual_reduction / 100)
                support_price = base_price * 0.20 * first_product.support_years
                template_data["supportprice"] = f"{support_price:,.2f}"
            else:
                template_data["supportprice"] = "0"
        
        # Traiter le document
        process_word_document(doc, template_data, product_dict if first_product else None)
        
        # Sauvegarder en mémoire
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{request.proposal_number}.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
            }
        )
    except Exception as e:
        logger.error(f"Erreur: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-invoice-excel")
async def generate_invoice_excel(request: dict):
    """Génère un Excel d'invoice à partir des données en utilisant le template Test.xlsx"""
    try:
        logger.info(f"Génération Excel Invoice pour: {request['invoice']['invoice_number']}")
        
        # Importer les modules nécessaires
        from openpyxl import load_workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        from openpyxl.utils import get_column_letter
        from datetime import datetime
        from io import BytesIO
        
        # Chemin vers le template Excel
        template_path = "/Users/alexiszebidi/Document/ATDI/CRM/Application-taury/sidecar-python/Test.xlsx"
        
        # Charger le template Excel
        wb = load_workbook(template_path)
        ws = wb.active
        
        # Extraire les données
        invoice = request['invoice']
        company = request['company']
        contact = request.get('contact')
        products = request['products']
        
        modifications = []
        
        # C13 - Nom de la compagnie
        if company.get('name'):
            ws['C13'] = company['name']
            modifications.append(f"C13: {company['name']}")
            logger.info(f"Cellule C13 modifiée: {company['name']}")
        
        # C6 - Purchase Order (avec le numéro de commande)
        purchase_order = invoice.get('purchase_order', '')
        if purchase_order:
            ws['C6'] = f"Purchase Order: {purchase_order}"
            modifications.append(f"C6: Purchase Order: {purchase_order}")
            logger.info(f"Cellule C6 modifiée: Purchase Order: {purchase_order}")
        else:
            ws['C6'] = "Purchase Order: Not specified"
            modifications.append(f"C6: Purchase Order: Not specified")
            logger.info(f"Cellule C6 modifiée: Purchase Order: Not specified")
        
        # E3 - Numéro de facture
        if invoice.get('invoice_number'):
            ws['E3'] = invoice['invoice_number']
            modifications.append(f"E3: {invoice['invoice_number']}")
            logger.info(f"Cellule E3 modifiée: {invoice['invoice_number']}")
        
        # F3 - Date du jour
        today = datetime.now().strftime("%d/%m/%Y")
        ws['F3'] = today
        modifications.append(f"F3: {today}")
        logger.info(f"Cellule F3 modifiée: {today}")
        
        # C8 - Purchase order date (utilise l'issue_date de la facture)
        issue_date = invoice.get('issue_date', '')
        if issue_date:
            try:
                if isinstance(issue_date, str):
                    parsed_date = datetime.strptime(issue_date, "%Y-%m-%d")
                    formatted_date = parsed_date.strftime("%d %B %Y")
                else:
                    formatted_date = issue_date.strftime("%d %B %Y")
                
                ws['C8'] = f"Purchase order date: {formatted_date}"
                modifications.append(f"C8: Purchase order date: {formatted_date}")
                logger.info(f"Cellule C8 modifiée: Purchase order date: {formatted_date}")
            except Exception as e:
                logger.warning(f"Erreur formatage date C8: {e}")
                ws['C8'] = "Purchase order date: Not specified"
                modifications.append(f"C8: Purchase order date: Not specified")
        else:
            ws['C8'] = "Purchase order date: Not specified"
            modifications.append(f"C8: Purchase order date: Not specified")
            logger.info(f"Cellule C8 modifiée: Purchase order date: Not specified")
        
        # C10 - Commercial in charge (avec valeur par défaut si vide)
        commercial_in_charge = invoice.get('commercial_in_charge', '')
        if commercial_in_charge:
            ws['C10'] = f"Commercial in charge: {commercial_in_charge}"
            modifications.append(f"C10: Commercial in charge: {commercial_in_charge}")
            logger.info(f"Cellule C10 modifiée: Commercial in charge: {commercial_in_charge}")
        else:
            ws['C10'] = "Commercial in charge: Not specified"
            modifications.append(f"C10: Commercial in charge: Not specified")
            logger.info(f"Cellule C10 modifiée: Commercial in charge: Not specified")
        
        # Ajouter les produits dans les lignes (commencer à la ligne 41)
        current_row = 41
        item_number = 1
        total_general = 0
        
        if not products:
            # Si pas de produits détaillés, créer un produit simple
            if invoice.get('total_amount'):
                total_amount = float(invoice.get('total_amount', 0))
                
                # Ligne A - Numéro d'item
                ws[f'A{current_row}'] = item_number
                modifications.append(f"A{current_row}: {item_number}")
                
                # Ligne B - Nom du service (avec quantité)
                service_name = f"Service pour {company.get('name', 'Client')}"
                ws[f'B{current_row}'] = service_name
                modifications.append(f"B{current_row}: {service_name}")
                
                # Ligne D - Quantité (1 par défaut)
                ws[f'D{current_row}'] = 1
                modifications.append(f"D{current_row}: 1")
                
                # Ligne E - Prix unitaire
                ws[f'E{current_row}'] = total_amount
                modifications.append(f"E{current_row}: {total_amount}")
                
                # Ligne F - Prix total
                ws[f'F{current_row}'] = total_amount
                modifications.append(f"F{current_row}: {total_amount}")
                
                total_general += total_amount
                
                # Formater les cellules de prix
                try:
                    ws[f'E{current_row}'].number_format = '#,##0.00 €'
                    ws[f'F{current_row}'].number_format = '#,##0.00 €'
                    modifications.append(f"E{current_row},F{current_row}: formaté en euros")
                except Exception as e:
                    logger.warning(f"Erreur formatage prix: {e}")
        else:
            # Traiter chaque produit avec services séparés (logique de l'ancien api.py)
            for product in products:
                product_name = product.get('product_type', 'Unknown Product')
                has_licence = product.get('licence', False)
                has_training = product.get('training', False)
                has_support = product.get('support', False)
                
                # Déterminer l'ordre des services
                services = []
                if has_licence:
                    services.append('licence')
                if has_training:
                    services.append('training')
                if has_support:
                    services.append('support')
                
                # Traiter chaque service sur une ligne séparée
                for service_type in services:
                    row = current_row
                    
                    # Colonne A - Numéro d'item
                    ws[f'A{row}'] = item_number
                    modifications.append(f"A{row}: {item_number}")
                    
                    # Colonne B - Nom du service (toujours complet)
                    if service_type == 'licence':
                        service_name = f"{product_name} Licence"
                    elif service_type == 'training':
                        service_name = f"{product_name} Training"
                    elif service_type == 'support':
                        service_name = f"{product_name} Support"
                    
                    ws[f'B{row}'] = service_name
                    modifications.append(f"B{row}: {service_name}")
                    
                    # Incrémenter le numéro d'item
                    item_number += 1
                    
                    # Colonne D - Quantité selon le service
                    quantity = 0
                    total_price = 0
                    
                    if service_type == 'licence':
                        # Quantité totale (nombre d'users)
                        quantity = product.get('user_count', 0) or 0
                        ws[f'D{row}'] = quantity
                        modifications.append(f"D{row}: {quantity}")
                    elif service_type == 'training':
                        # Nombre de jours de training
                        quantity = product.get('training_days', 0) or 0
                        ws[f'D{row}'] = quantity
                        modifications.append(f"D{row}: {quantity}")
                    elif service_type == 'support':
                        quantity = product.get('support_years', 0) or 0
                        ws[f'D{row}'] = quantity
                        modifications.append(f"D{row}: {quantity} (années de support)")
                    
                    # Colonne F - Prix selon le service
                    if service_type == 'licence':
                        # Calculer le prix licence
                        total_price_product = float(product.get('unit_price', 0) or 0)
                        annual_reduction = float(product.get('annual_reduction', 0) or 0)
                        users = int(quantity or 0)
                        
                        # Le prix total est déjà donné, on l'utilise directement
                        licence_price = total_price_product
                        
                        # Appliquer la réduction si présente
                        if annual_reduction > 0:
                            licence_price = licence_price * (100 - annual_reduction) / 100
                        
                        ws[f'F{row}'] = licence_price
                        modifications.append(f"F{row}: {licence_price}")
                        total_price += licence_price
                        total_general += licence_price
                        
                    elif service_type == 'training':
                        # Prix training (jours × prix par jour)
                        training_days = product.get('training_days', 0) or 0
                        training_cost_per_day = float(product.get('training_cost_per_day', 0) or 0)
                        training_price = training_days * training_cost_per_day
                        ws[f'F{row}'] = training_price
                        modifications.append(f"F{row}: {training_price}")
                        total_price += training_price
                        total_general += training_price
                        
                    elif service_type == 'support':
                        # Prix support (20% du prix licence)
                        unit_price = float(product.get('unit_price', 0) or 0)
                        annual_reduction = float(product.get('annual_reduction', 0) or 0)
                        users = int(quantity or 0)
                        
                        base_price = unit_price * users if users > 0 else unit_price
                        support_price = base_price * 0.2
                        
                        # Appliquer la réduction si présente
                        if annual_reduction > 0:
                            support_price = support_price * (100 - annual_reduction) / 100
                        
                        ws[f'F{row}'] = support_price
                        modifications.append(f"F{row}: {support_price}")
                        total_price += support_price
                        total_general += support_price
                    
                    # Colonne E - Prix unitaire (calculé automatiquement: F/D)
                    if quantity > 0:
                        unit_price = total_price / quantity
                        ws[f'E{row}'] = unit_price
                        modifications.append(f"E{row}: {unit_price} (prix unitaire: {total_price}/{quantity})")
                    else:
                        ws[f'E{row}'] = 0
                        modifications.append(f"E{row}: 0")
                    
                    # Formater les cellules de prix
                    try:
                        ws[f'E{row}'].number_format = '#,##0.00 €'
                        ws[f'F{row}'].number_format = '#,##0.00 €'
                        modifications.append(f"E{row},F{row}: formaté en euros")
                    except Exception as e:
                        logger.warning(f"Erreur formatage prix E{row},F{row}: {e}")
                    
                    # Passer à la ligne suivante avec un espace
                    current_row += 2  # +2 pour créer un espace entre chaque item
        
        logger.info(f"Total modifications Excel: {len(modifications)}")
        for mod in modifications:
            logger.info(f"  - {mod}")
        
        # Sauvegarder en mémoire
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        filename = f"Invoice_{invoice['invoice_number']}.xlsx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
            }
        )
    except Exception as e:
        logger.error(f"Erreur génération Excel: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health():
    """Health check"""
    return {"status": "ok", "version": "1.0.0"}

if __name__ == "__main__":
    logger.info(f"🚀 Document Generator sur {API_HOST}:{API_PORT}")
    logger.info(f"📄 Template: {TEMPLATE_PATH}")
    uvicorn.run(app, host=API_HOST, port=API_PORT, log_level="error")

